from netmiko import ConnectHandler
import getpass
import sys
from device import *
from docx import Document
from excel_commands import commands
import re

if __name__ == "__main__":

	# def type_of_config(command, commands):

	# 	return commands.get(command)

	# print ("Script for SSH to device, Please enter your credential")
	# device['username']=input("User name: ")
	# device['password']=getpass.getpass()

	#Establishing SSH connection
	switch = ConnectHandler(**device)
	print('SSH connection established')
	#changing to enable mode
	switch.enable()

	print('enabled')

	switch.send_config_set(commands)

	doc_text = str(switch.send_command('sho run | sec interface'))

	doc_format_text = re.split(r'(interface \w+/\w/\d+)', doc_text)

	document = Document()
	document.add_heading(doc_format_text[1])
	document.add_paragraph(doc_format_text[2])

	for i in range(3,len(doc_format_text)):
		if i%2 ==1: document.add_heading(doc_format_text[i])
		else: document.add_paragraph(doc_format_text[i])


	print(switch.send_command('sho run | sec GigabitEthernet1/0'))

	print('\n'*2 + 'saving G1config as a word document\n')
	doc_name = input('Enter the name for the document:')
	doc_name += '.docx'
	document.save('./config_files/' + doc_name)

	print('commands done, disconnecting...')

	switch.disconnect()
	print('Connection ended')